import streamlit as st

st.set_page_config(
    page_title="四年精讀聖經",
    page_icon="icon.jpg",
    layout="wide"
)

# 使用更強效的新版隱藏語法
hide_streamlit_style = """
<style>
/* 隱藏右上角主選單、頁尾與頂部導覽列 */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
[data-testid="stHeader"] {visibility: hidden; display: none;}
[data-testid="stToolbar"] {visibility: hidden; display: none;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

import pandas as pd
import datetime
import os
import logging
import io
import docx
import gspread
from google.oauth2.service_account import Credentials
import streamlit.components.v1 as components
from googleapiclient.discovery import build

# ==========================================
# 簽到二次確認彈窗
# ==========================================
@st.dialog("簽到確認")
def confirm_checkin_dialog(member_name, week_display, week_key, missing_weeks):
    st.markdown(f"👉 確定要為 **{member_name}** 辦理 **{week_display}** 的簽到嗎？")
    
    if missing_weeks:
        st.info(f"💡 系統將一併自動為您補簽過往未簽到的 **{len(missing_weeks)}** 週進度！")
        
    col1, col2 = st.columns(2)
    with col1:
        if st.button("✅ 確定簽到", type="primary", use_container_width=True):
            records_to_add = [(week_key, member_name)]
            for m_item in missing_weeks:
                records_to_add.append((m_item["key"], member_name))
            
            add_batch_records(records_to_add)
            
            if missing_weeks:
                st.toast(f"🎉 簽到成功！已一併補齊過往 {len(missing_weeks)} 週進度！")
            else:
                st.toast("🎉 簽到成功！")
                
            st.session_state.scroll_target = "divider-top-anchor"
            st.rerun()
            
    with col2:
        if st.button("❌ 取消", type="secondary", use_container_width=True):
            st.rerun()

# 設定 Logging 紀錄
logging.basicConfig(level=logging.INFO)

# ==========================================
# 1. 基礎設定與常數
# ==========================================
MEMBERS_FILE = "church_members.csv"
VERSES_FILE = "verses.csv"
ATTENDANCE_FILE = "attendance_records.csv"
GUIDE_FOLDER_ID = "1-RkVxCZy9wS_2X6Huw5p2mWhv1b6l0HM"

ADMIN_PASSWORD = st.secrets.get("admin_password", "11190928")
PLAN_YEAR = 2

INITIAL_MEMBERS = [
    "周寶燕", "曾笑", "黃然玉", "吳妃玉", "楊游美麗", 
    "翁淑美", "石美莎", "單麗蘭", "鄭富美", "李鶯芳", 
    "趙文崇", "李應昌", "賴健文", "林春妙", "邱文雀", 
    "梁垠盤", "陳宜宏", "郭彩梅", "林春桃", "鳳姐", 
    "黃敏生", "吳秀卉", "陳安俐", "程乃珍", "蕭慧麗", 
    "蔡慧俐", "林雅谷", "李俊修", "林淑惠", "盧正亮", 
    "翁春祝", "劉淑珠", "葉雅雲", "林雅音", "趙文川",
    "邱聖富"
]

st.set_page_config(page_title="四年精讀聖經運動簽到系統", page_icon="📖", layout="wide")

# ==========================================
# 2. 輔助與 GCP 憑證函式
# ==========================================
def get_gcp_credentials():
    if "gcp_service_account" not in st.secrets:
        return None
    scope = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive"
    ]
    creds_dict = dict(st.secrets["gcp_service_account"])
    if "private_key" in creds_dict:
        pk = creds_dict["private_key"].replace("\\n", "\n")
        if not pk.startswith("-----BEGIN PRIVATE KEY-----"):
            pk = "-----BEGIN PRIVATE KEY-----\n" + pk
        if not pk.endswith("-----END PRIVATE KEY-----"):
            pk = pk + "\n-----END PRIVATE KEY-----"
        creds_dict["private_key"] = pk.strip()

    return Credentials.from_service_account_info(creds_dict, scopes=scope)

@st.cache_resource
def get_drive_service():
    creds = get_gcp_credentials()
    if not creds:
        return None
    scopes = ["https://www.googleapis.com/auth/drive.readonly"]
    scoped_creds = creds.with_scopes(scopes)
    return build("drive", "v3", credentials=scoped_creds)

@st.cache_data(ttl=60)
def fetch_docx_content(week_num, target_date=None):
    try:
        service = get_drive_service()
        if not service:
            return None
        clean_week = "".join(filter(str.isdigit, str(week_num)))
        query = f"'{GUIDE_FOLDER_ID}' in parents and name contains '{selected_year}' and name contains '{clean_week}' and trashed = false"
        results = service.files().list(q=query, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True).execute()
        files = results.get("files", [])
        
        target_file = files[0] if files else None
        if not target_file:
            return "找不到檔案"

        request = service.files().get_media(fileId=target_file["id"])
        file_bytes = io.BytesIO(request.execute())
        
        doc = docx.Document(file_bytes)
        
        if not target_date:
            import re
            full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
    
            full_text = re.sub(r'\[.*?\]', '', full_text)
            full_text = re.sub(r'［.*?］', '', full_text)
    
            return full_text
            
        extracted_lines = []
        is_recording = False
        
        all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip() != ""]
        clean_target = str(target_date).replace(" ", "")
        
        for text in all_paragraphs:
            clean_line = text.replace(" ", "")
            
            if f"[DATE:{clean_target}]" in clean_line:
                is_recording = True
                continue
                
            if is_recording:
                if "[END_DATE]" in clean_line or ("[DATE:" in clean_line and clean_target not in clean_line):
                    break
                extracted_lines.append(text)
                
        if extracted_lines:
            return "\n\n".join(extracted_lines).strip()
        else:
            return f"⚠️ 找不到對應 `{target_date}` 的範圍。檔案內的前幾行標記範例：\n\n" + "\n---\n".join(all_paragraphs[:5])
        
    except Exception as e:
        return f"⚠️ 發生錯誤：{e}"

# ==========================================
# 3. Google Drive 動態抓取圖片網址 (帶年份)
# ==========================================
@st.cache_data(ttl=300)
def get_gdrive_image_url(year_num, week_num):
    try:
        creds = get_gcp_credentials()
        if not creds:
            return None
        
        drive_service = build('drive', 'v3', credentials=creds)

        folder_id = st.secrets.get("drive_folder_id", None)
        if not folder_id:
            return None
            
        if "folders/" in folder_id:
            folder_id = folder_id.split("folders/")[1].split("?")[0]

        actual_year = 2026 - (PLAN_YEAR - year_num)
        search_year = str(actual_year)
        search_term_1 = f"第{week_num}周"
        search_term_2 = f"第{week_num}週"

        query = (
            f"'{folder_id}' in parents and name contains '{search_year}' "
            f"and (name contains '{search_term_1}' or name contains '{search_term_2}') "
            f"and trashed = false"
        )
        
        results = drive_service.files().list(
            q=query, 
            fields="files(id, name)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()

        files = results.get('files', [])
        if files:
            file_id = files[0]['id']
            return f"https://lh3.googleusercontent.com/d/{file_id}"
            
    except Exception as e:
        logging.error(f"從 Google Drive 搜尋 Y{year_num}-W{week_num} 圖片失敗: {e}")
    
    return None

# ==========================================
# 4. 資料庫與簽到邏輯
# ==========================================
def load_attendance():
    try:
        creds = get_gcp_credentials()
        if creds:
            client = gspread.authorize(creds)
            sheet_name = st.secrets.get("spreadsheet_name", "Church_Attendance")
            sheet = client.open(sheet_name).sheet1
            
            rows = sheet.get_all_records()
            if rows:
                df_gs = pd.DataFrame(rows)
                expected_cols = ["week_key", "member_name", "timestamp"]
                for col in expected_cols:
                    if col not in df_gs.columns:
                        df_gs[col] = ""
                    else:
                        df_gs[col] = df_gs[col].astype(str).str.strip()
                
                df_gs[expected_cols].to_csv(ATTENDANCE_FILE, index=False, encoding="utf-8-sig")
                return df_gs[expected_cols]
    except Exception as e:
        logging.error(f"從 Google Sheets 讀取簽到紀錄失敗，改用本機快取: {e}")

    if os.path.exists(ATTENDANCE_FILE):
        try:
            df = pd.read_csv(ATTENDANCE_FILE, dtype=str)
            for col in ["week_key", "member_name", "timestamp"]:
                if col not in df.columns:
                    df[col] = ""
                else:
                    df[col] = df[col].astype(str).str.strip()
            return df
        except Exception as e:
            logging.error(f"讀取本機簽到紀錄失敗: {e}")
            
    return pd.DataFrame(columns=["week_key", "member_name", "timestamp"])

def save_attendance(df):
    df.to_csv(ATTENDANCE_FILE, index=False, encoding="utf-8-sig")

def sync_to_gsheet_async(new_rows_list):
    try:
        creds = get_gcp_credentials()
        if not creds:
            return
        client = gspread.authorize(creds)
        sheet_name = st.secrets.get("spreadsheet_name", "Church_Attendance")
        sheet = client.open(sheet_name).sheet1
        sheet.append_rows(new_rows_list)
    except Exception as e:
        logging.error(f"Google Sheets 同步失敗: {e}")

def add_batch_records(records_list):
    df = load_attendance()
    now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    new_rows = []
    
    for week_key, member_name in records_list:
        week_key = str(week_key).strip()
        member_name = str(member_name).strip()
        match = df[(df["week_key"] == week_key) & (df["member_name"] == member_name)]
        if match.empty:
            new_rows.append({"week_key": week_key, "member_name": member_name, "timestamp": now_str})
            
    if new_rows:
        new_df = pd.DataFrame(new_rows)
        df = pd.concat([df, new_df], ignore_index=True)
        save_attendance(df)
        gsheet_rows = [[r["week_key"], r["member_name"], r["timestamp"]] for r in new_rows]
        sync_to_gsheet_async(gsheet_rows)

def delete_single_record(week_key, member_name):
    df = load_attendance()
    week_key = str(week_key).strip()
    member_name = str(member_name).strip()
    df_new = df[~((df["week_key"] == week_key) & (df["member_name"] == member_name))]
    save_attendance(df_new)
    return True

def load_members():
    if os.path.exists(MEMBERS_FILE):
        try:
            df_m = pd.read_csv(MEMBERS_FILE, encoding="utf-8-sig")
            if "member_name" in df_m.columns and not df_m.empty:
                df_m["member_name"] = df_m["member_name"].astype(str).str.strip()
                return df_m
        except Exception as e:
            logging.error(f"讀取會友名單失敗: {e}")

    df_m = pd.DataFrame({"member_name": INITIAL_MEMBERS})
    df_m.to_csv(MEMBERS_FILE, index=False, encoding="utf-8-sig")
    return df_m

def save_members(members_list):
    pd.DataFrame({"member_name": members_list}).to_csv(MEMBERS_FILE, index=False, encoding="utf-8-sig")

def get_weekly_verse(week_num):
    fallback = {
        "verse": "「你的話是我腳前的燈，是我路上的光。」", 
        "ref": "詩篇 119:105",
        "encouragement": "讓上帝的話語成為你每日的亮光與引導！"
    }
    if os.path.exists(VERSES_FILE):
        try:
            v_df = pd.read_csv(VERSES_FILE)
            if not v_df.empty:
                row = v_df.iloc[(week_num - 1) % len(v_df)]
                return {
                    "verse": str(row["verse"]), 
                    "ref": str(row["ref"]),
                    "encouragement": str(row.get("encouragement", ""))
                }
        except Exception:
            pass
    return fallback

def get_current_year_and_week():
    today = datetime.date.today()
    
    # 關鍵技巧：把今天往後加 3 天
    # 這樣只要今天是週五、週六、週日，或是隔週一到週四，
    # 透過 ISO 週數計算時，就會剛好落在「新的一週」裡！
    adjusted_today = today + datetime.timedelta(days=3)
    
    iso_year, iso_week, _ = adjusted_today.isocalendar()
    
    # 設定您的計畫年份與調整後的週數
    current_year = 2
    current_week = iso_week  # 今天（星期五）加上 3 天後剛好會算出 37 週！
    
    return current_year, current_week
    
def generate_pivot_report(target_year, max_week):
    df_att = load_attendance()
    members = load_members()["member_name"].tolist()

    report_data = []
    week_cols = [f"Y{target_year}-W{w:02d}" for w in range(1, max_week + 1)]

    for m in members:
        row = {"member_name": m}
        completed_count = 0
        m_signed = set(df_att[df_att["member_name"] == m]["week_key"].tolist())

        for w_col in week_cols:
            if w_col in m_signed:
                row[w_col] = "⚪ 已讀"
                completed_count += 1
            else:
                row[w_col] = "❌"

        row["完成週數"] = f"{completed_count} / {max_week}"
        row["完成率"] = f"{(completed_count / max_week * 100):.2f}%" if max_week > 0 else "0.00%"
        report_data.append(row)

    df_report = pd.DataFrame(report_data)
    cols_order = ["member_name", "完成週數", "完成率"] + week_cols
    return df_report[cols_order]

# 取得當前的年份與週次
PLAN_YEAR, current_week_num = get_current_year_and_week()

# 顯示在畫面上時，就會完美呈現您要的格式：
# 例如：畫面標題自動顯示為 「最新讀經進度表 (第 2 年 - 第 37 週)」

# ==========================================
# 5. CSS 樣式
# ==========================================
st.markdown("""
    <style>
    html, body { max-width: 100vw; overflow-x: hidden; }
    h1 { font-size: clamp(26px, 6vw, 38px) !important; line-height: 1.3 !important; }

    div[data-baseweb="tab-list"] {
        gap: 10px !important;
        margin-bottom: 20px !important;
    }

    button[data-baseweb="tab"] {
        border-radius: 12px !important;
        padding: 12px 20px !important;
        margin: 2px !important;
        transition: all 0.2s ease-in-out !important;
        box-shadow: 0px 2px 5px rgba(0,0,0,0.08) !important;
    }

    div[data-testid="stTabs"] [role="tab"] p, 
    div[data-testid="stTabs"] [role="tab"] div {
        font-size: clamp(20px, 4.5vw, 24px) !important;
        font-weight: 900 !important;
        letter-spacing: 1px !important;
        line-height: 1.3 !important;
    }

    /* Tab 1: 會友簽到專區 */
    button[data-baseweb="tab"]:nth-child(1) { background-color: #ECFDF5 !important; border: 2.5px solid #10B981 !important; }
    button[data-baseweb="tab"]:nth-child(1) p { color: #047857 !important; }
    button[data-baseweb="tab"]:nth-child(1)[aria-selected="true"] { background-color: #059669 !important; border-color: #047857 !important; }
    button[data-baseweb="tab"]:nth-child(1)[aria-selected="true"] p { color: #FFFFFF !important; }

    /* Tab 2: 歷史讀經與導讀 */
    button[data-baseweb="tab"]:nth-child(2) { background-color: #EFF6FF !important; border: 2.5px solid #3B82F6 !important; }
    button[data-baseweb="tab"]:nth-child(2) p { color: #1D4ED8 !important; }
    button[data-baseweb="tab"]:nth-child(2)[aria-selected="true"] { background-color: #2563EB !important; border-color: #1D4ED8 !important; }
    button[data-baseweb="tab"]:nth-child(2)[aria-selected="true"] p { color: #FFFFFF !important; }

    /* 頁籤 3: 長者輔助資源 */
    button[data-baseweb="tab"]:nth-child(3) { background-color: #FAF5FF !important; border: 2.5px solid #8B5CF6 !important; }
    button[data-baseweb="tab"]:nth-child(3) p { color: #6D28D9 !important; }
    button[data-baseweb="tab"]:nth-child(3)[aria-selected="true"] { background-color: #7C3AED !important; border-color: #6D28D9 !important; }
    button[data-baseweb="tab"]:nth-child(3)[aria-selected="true"] p { color: #FFFFFF !important; }
    
    /* Tab 4: 後台統計管理 */
    button[data-baseweb="tab"]:nth-child(4) { background-color: #F8FAFC !important; border: 2.5px solid #64748B !important; }
    button[data-baseweb="tab"]:nth-child(4) p { color: #334155 !important; }
    button[data-baseweb="tab"]:nth-child(4)[aria-selected="true"] { background-color: #475569 !important; border-color: #334155 !important; }
    button[data-baseweb="tab"]:nth-child(4)[aria-selected="true"] p { color: #FFFFFF !important; }

    div[data-baseweb="tab-highlight"] { display: none !important; }

    div[data-aria-expanded] p, div[data-testid="stExpander"] summary p {
        font-size: clamp(20px, 4.8vw, 26px) !important;
        font-weight: 800 !important;
        line-height: 1.5 !important;
        color: #1E293B !important;
    }

    div[data-testid="stButton"] button {
        width: 100% !important;
        white-space: normal !important;
        word-break: break-word !important;
    }
    div[data-testid="stButton"] button p {
        font-size: clamp(20px, 5.5vw, 28px) !important;
        font-weight: 800 !important;
    }
    div[data-testid="stButton"] button[kind="secondary"] {
        min-height: 3.2em !important;
        padding: 10px 8px !important;
        border-radius: 14px !important;
        border: 2.5px solid #0284C7 !important;
        background-color: #FFFFFF !important;
        color: #0F172A !important;
        margin-bottom: 10px !important;
    }
    div[data-testid="stButton"] button[kind="secondary"]:hover { background-color: #E0F2FE !important; }
    div[data-testid="stButton"] button[kind="primary"] {
        min-height: 3.5em !important;
        border-radius: 14px !important;
        background-color: #059669 !important;
        margin-bottom: 10px !important;
    }
    div[data-testid="stButton"] button[kind="primary"] p { color: #FFFFFF !important; }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 6. 主介面
# ==========================================
if "current_member" not in st.session_state:
    st.session_state.current_member = None

# 確保呼叫時都是這樣寫：
PLAN_YEAR, current_week_num = get_current_year_and_week()

current_week_key = f"Y{PLAN_YEAR}-W{current_week_num:02d}"
current_week_display = f"第 {PLAN_YEAR} 年 - 第 {current_week_num:02d} 週"

df_members = load_members()
member_list = df_members["member_name"].tolist()
df_attendance = load_attendance()

st.title(f"📖 最新讀經進度表（{current_week_display}）")

# 嚴格確保第 1、2 頁籤不變，第 3 頁為雲端資料，第 4 頁為後台
tab_user, tab_history, tab_resource, tab_admin = st.tabs([
    "✍️ 會友簽到專區", 
    "🗓️ 讀經暨導讀查詢系統", 
    "📖 讀經資源專區",
    "🔒 後台統計管理"
])

# ------------------------------------------
# TAB 1: 會友簽到專區 (維持原樣)
# ------------------------------------------
with tab_user:
    current_img_url = get_gdrive_image_url(PLAN_YEAR, current_week_num)
    if current_img_url:
        st.image(current_img_url, caption=f"【最新進度】{current_week_display}", use_container_width=True)
    else:
        st.info(f"📌 目前為【{current_week_display}】簽到（雲端硬碟尚未找到第 {current_week_num} 週進度表）。")

    st.markdown("<div id='divider-top-anchor'></div>", unsafe_allow_html=True)
    st.divider()

    if "scroll_target" not in st.session_state:
        st.session_state.scroll_target = None
    if "open_section" not in st.session_state:
        st.session_state.open_section = None

    if st.session_state.scroll_target:
        target_id = st.session_state.scroll_target
        components.html(
            f"""
            <script>
                setTimeout(function() {{
                    var el = window.parent.document.getElementById('{target_id}');
                    if (el) {{
                        el.scrollIntoView({{behavior: 'smooth', block: 'start'}});
                    }}
                }}, 150);
            </script>
            """,
            height=0,
        )
        st.session_state.scroll_target = None

    if st.session_state.current_member is None:
        st.markdown("<div id='members-list-top'></div>", unsafe_allow_html=True)
        st.markdown("### 👇 請點擊您所屬的分區展開名字列表：")

        valid_members = [
            m for m in member_list 
            if m and str(m).strip() and not str(m).startswith("會友")
        ]

        chunk_size = 10
        total_valid = len(valid_members)

        for i in range(0, total_valid, chunk_size):
            chunk = valid_members[i:i + chunk_size]
            page_num = (i // chunk_size) + 1

            if chunk:
                names_text = "、".join(chunk)
                is_this_open = (st.session_state.open_section == page_num)

                toggle_icon = "🔽" if is_this_open else "▶️"
                header_label = f"{toggle_icon} 📦 【第 {page_num} 區】 {names_text}"

                if st.button(header_label, key=f"sec_toggle_{page_num}", type="secondary", use_container_width=True):
                    if is_this_open:
                        st.session_state.open_section = None
                    else:
                        st.session_state.open_section = page_num
                        st.session_state.scroll_target = f"line-anchor-{page_num}"
                    st.rerun()

                if is_this_open:
                    st.markdown(f"<div id='line-anchor-{page_num}'></div>", unsafe_allow_html=True)
                    st.divider()

                    col1, col2 = st.columns(2)
                    mid = (len(chunk) + 1) // 2

                    with col1:
                        for name in chunk[:mid]:
                            is_signed = not df_attendance[(df_attendance["week_key"] == current_week_key) & (df_attendance["member_name"] == name)].empty
                            status_icon = "✅" if is_signed else "👤"
                            if st.button(f"{status_icon} {name}", key=f"btn_dyn_{page_num}_{name}_c1", type="secondary", use_container_width=True):
                                st.session_state.current_member = name
                                st.session_state.scroll_target = "divider-top-anchor"
                                st.rerun()

                    with col2:
                        for name in chunk[mid:]:
                            is_signed = not df_attendance[(df_attendance["week_key"] == current_week_key) & (df_attendance["member_name"] == name)].empty
                            status_icon = "✅" if is_signed else "👤"
                            if st.button(f"{status_icon} {name}", key=f"btn_dyn_{page_num}_{name}_c2", type="secondary", use_container_width=True):
                                st.session_state.current_member = name
                                st.session_state.scroll_target = "divider-top-anchor"
                                st.rerun()

    else:
        member_name = st.session_state.current_member

        if st.button("⬅️ 返回選擇名字列表", type="secondary", use_container_width=True):
            st.session_state.current_member = None
            current_sec = st.session_state.open_section
            if current_sec:
                st.session_state.scroll_target = f"line-anchor-{current_sec}"
            else:
                st.session_state.scroll_target = "members-list-top"
            st.rerun()

        st.markdown(f"## 👤 {member_name} 的讀經專頁")
        
        signed_weeks = df_attendance[df_attendance["member_name"] == member_name]["week_key"].tolist()
        
        missing_weeks_info = []
        for w in range(1, current_week_num):
            w_key = f"Y{PLAN_YEAR}-W{w:02d}"
            w_display = f"第 {PLAN_YEAR} 年 - 第 {w:02d} 週"
            if w_key not in signed_weeks:
                missing_weeks_info.append({"key": w_key, "display": w_display, "week_num": w})

        st.markdown(f"### 📍 【本週進度】{current_week_display}")

        is_signed = not df_attendance[(df_attendance["week_key"] == current_week_key) & (df_attendance["member_name"] == member_name)].empty

        if is_signed:
            st.success(f"🎉 **{member_name}**，您已完成本週讀經進度，願主保守力上加力恩上加恩！")
        else:
            if st.button(f"🟢 若完成【{current_week_display}】請按此簽到", type="primary", use_container_width=True):
                confirm_checkin_dialog(member_name, current_week_display, current_week_key, missing_weeks_info)
                
                records_to_add = [(current_week_key, member_name)]
                for m_item in missing_weeks_info:
                    records_to_add.append((m_item["key"], member_name))
                
                add_batch_records(records_to_add)

                if missing_weeks_info:
                    st.toast(f"🎉 簽到成功！並已自動為您補齊過往 {len(missing_weeks_info)} 週進度！")
                else:
                    st.toast("🎉 簽到成功！")
                
                st.session_state.scroll_target = "divider-top-anchor"
                st.rerun()

        st.divider()
        st.markdown("### 🟡 【過往進度補簽狀態】")

        if missing_weeks_info:
            if not is_signed:
                st.warning(f"⚠️ 您尚有 **{len(missing_weeks_info)}** 週過往進度尚未簽到，點擊以下按鈕可單獨補簽：")
            else:
                st.info(f"📌 您先前尚有 **{len(missing_weeks_info)}** 週紀錄未補齊，可點擊下方按鈕單獨補簽：")

            mid_m = (len(missing_weeks_info) + 1) // 2

            mc1, mc2 = st.columns(2)
            with mc1:
                for item in missing_weeks_info[:mid_m]:
                    if st.button(f"🟡 {item['display']}", key=f"miss_{member_name}_{item['key']}_c1", type="secondary", use_container_width=True):
                        add_batch_records([(item["key"], member_name)])
                        st.toast(f"✅ 已成功補簽 `{item['display']}`！")
                        st.session_state.scroll_target = "divider-top-anchor"
                        st.rerun()
            with mc2:
                for item in missing_weeks_info[mid_m:]:
                    if st.button(f"🟡 {item['display']}", key=f"miss_{member_name}_{item['key']}_c2", type="secondary", use_container_width=True):
                        add_batch_records([(item["key"], member_name)])
                        st.toast(f"✅ 已成功補簽 `{item['display']}`！")
                        st.session_state.scroll_target = "divider-top-anchor"
                        st.rerun()
        else:
            st.success("🎉 過往進度已全部完成，無需補簽！")

    st.divider()
    verse_info = get_weekly_verse(current_week_num)
    st.markdown(f"📖 **本週靈修經文**：{verse_info['ref']}")
    st.markdown(f"> *{verse_info['verse']}*")
    if verse_info.get('encouragement'):
        st.markdown(f"💬 **心靈補給**：{verse_info['encouragement']}")

# ------------------------------------------
# TAB 2: 歷史讀經與導讀查詢 (維持原樣)
# ------------------------------------------
with tab_history:
    st.markdown("### 🗓️ 歷史讀經進度表與導讀查詢")

    col_y, col_w = st.columns([1, 2])
    with col_y:
        selected_year = st.selectbox("請選擇年份：", [f"第 {y} 年 (Y{y})" for y in range(PLAN_YEAR, 0, -1)], index=0)
        target_y_num = int(selected_year.split("第 ")[1].split(" 年")[0])

    with col_w:
        max_w_display = current_week_num if target_y_num == PLAN_YEAR else 52
        week_options = [f"第 {w:02d} 週" for w in range(max_w_display, 0, -1)]
        selected_w_label = st.selectbox("請選擇週數：", week_options, index=0)
        target_w_num = int(selected_w_label.replace("第 ", "").replace(" 週", ""))

    history_img_url = get_gdrive_image_url(target_y_num, target_w_num)

    if history_img_url:
        st.image(history_img_url, caption=f"【第 {target_y_num} 年 - 第 {target_w_num:02d} 週】進度對照表", use_container_width=True)
    else:
        st.warning(f"📌 雲端硬碟中尚未找到【第 {target_y_num} 年 - 第 {target_w_num:02d} 週】的進度表圖片。")

    st.divider()
    st.markdown(f"### 📖 第 {target_w_num} 週 導讀內容閱覽")

    view_mode = st.radio(
        "請選擇檢視模式：",
        ["📜 全文導讀", "📅 按天切換閱讀 (Day 1 - Day 7)"],
        horizontal=True
    )

    if view_mode == "📅 按天切換閱讀 (Day 1 - Day 7)":
        selected_day = st.selectbox(
            "選擇天數：",
            [f"第 {i} 天" for i in range(1, 8)]
        )
        st.caption("💡 提示：導讀 Word 檔為全週彙整，您也可以隨時切換回「全文導讀」使用滾輪流暢瀏覽。")
    else:
        selected_day = None

    with st.spinner("正在從雲端硬碟導讀資料夾抓取檔案中..."):
        doc_content = fetch_docx_content(target_w_num, target_date=selected_day)

    if not doc_content:
        st.info(f"💡 雲端硬碟導讀資料夾中尚未找到第 {target_w_num} 週的 Word 導讀檔案。")
    else:
        display_text = doc_content
        
        if view_mode == "📜 全文導讀":
            import re
            display_text = re.sub(r'\[.*?\]', '', display_text)

        st.markdown(
            f"""
            <div style="
                height: 450px; 
                overflow-y: scroll; 
                background-color: #f8f9fa; 
                padding: 20px; 
                border-radius: 10px; 
                border: 1px solid #cbd5e1;
                line-height: 1.8;
                font-size: 16px;
                color: #1e293b;
                white-space: pre-wrap;
                box-shadow: inset 0 2px 4px rgba(0,0,0,0.05);
            ">
                {display_text}
            </div>
            """,
            unsafe_allow_html=True
        )

# ------------------------------------------
# TAB 3: 長者輔助資源 (包含認識經卷與有聲導讀)
# ------------------------------------------
with tab_resource:
    st.markdown("### 🎧 長者讀經輔助資源（參考專區）")
    st.info("💡 這裡提供給長輩與弟兄姊妹作為輔助參考的聲音導讀、經卷介紹與操作提醒，點擊下方按鈕即可參考：")

    st.markdown("---")

    # 區塊 1：認識經卷圖框與解說
    st.markdown("#### 📚 認識聖經經卷與背景")
    st.markdown("幫助長輩在讀經前快速了解各卷書的作者、寫作背景與核心主題：")
    
    col_book1, col_book2 = st.columns(2)
    with col_book1:
        st.markdown(
            """
            <div style="background-color: #F8FAFC; padding: 15px; border-radius: 10px; border: 2px solid #3B82F6;">
                <b>📖 華人基督徒查經資料網</b><br>
                <p style="font-size: 14px; color: #4B5563; margin-top: 5px;">提供純正和周詳的查經資料。</p>
                <a href="https://www.ccbiblestudy.org/index-T.htm" target="_blank" style="font-weight: bold; color: #2563EB;">👉 各經卷拾惠、例證、註解</a>
            </div>
            """,
            unsafe_allow_html=True
        )
    with col_book2:
        st.markdown(
            """
            <div style="background-color: #F8FAFC; padding: 15px; border-radius: 10px; border: 2px solid #10B981;">
                <b>📘 認識聖經各經卷SoundOn</b><br>
                <p style="font-size: 14px; color: #4B5563; margin-top: 5px;">認識每卷書背景與主題，明白上帝的本質與作為。</p>
                <a href="https://player.soundon.fm/p/49f6e2a8-a4c8-463c-9c97-7d4e4a8a4188" target="_blank" style="font-weight: bold; color: #059669;">👉 認識聖經經卷系列，點擊聆聽</a>
            </div>
            """,
            unsafe_allow_html=True
        )

    st.markdown("---")

    # 區塊 2：聲音導讀資源
    st.markdown("#### 🎙️ 推薦有聲導讀 / Podcast 資源")
    st.markdown("若長輩看字較吃力，或是希望在休閒、散步時聆聽經文導讀，可參考以下頻道：")
    
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        st.markdown(
            """
            <div style="background-color: #F3F4F6; padding: 15px; border-radius: 10px; border-left: 5px solid #8B5CF6;">
                <b>🎧 SoundOn 經卷註解導讀PodCast頻道 </b><br>
                <p style="font-size: 14px; color: #4B5563; margin-top: 5px;">適合輕鬆聆聽各周進度經卷生活應用、屬靈省司。</p>
                <a href="https://player.soundon.fm/p/28cbcb5d-2a87-4bb8-8b89-a3c2ccae77f8" target="_blank" style="font-weight: bold; color: #7C3AED;">👉 點擊前往聆聽</a>
            </div>
            """,
            unsafe_allow_html=True
        )
    with col_r2:
        st.markdown(
            """
            <div style="background-color: #F3F4F6; padding: 15px; border-radius: 10px; border-left: 5px solid #8B5CF6;">
                <b>📖 進度表盧俊義牧師導讀 </b><br>
                <p style="font-size: 14px; color: #4B5563; margin-top: 5px;">幫助快速掌握每週讀經進度的核心信息。</p>
                <a href="https://player.soundon.fm/p/520fefe3-1e30-4024-bcb1-260d1594bdf7" target="_blank" style="font-weight: bold; color: #7C3AED;">👉 點擊前往聆聽</a>
            </div>
            """,
            unsafe_allow_html=True
        )

    st.markdown("---")

    # 區塊 3：實用好幫手與操作提醒
    st.markdown("#### 📱 長輩操作小撇步")
    st.markdown(
        """
        * **字體放大**：若手機字體太小，可利用手機螢幕的放大功能或直接在瀏覽器設定中放大頁面顯示比例。
        * **現場協助**：若長輩在操作上有任何困難，歡迎隨時在聚會時請同工協助引導！
        """
    )

# ------------------------------------------
# TAB 4: 後台統計與管理 (第四個頁籤)
# ------------------------------------------
with tab_admin:
    st.subheader("🔒 管理者控制台")
    pwd = st.text_input("請輸入管理者密碼：", type="password")

    if pwd == ADMIN_PASSWORD:
        st.success("🔓 驗證成功，歡迎進入後台管理系統！")

        admin_sub_tab1, admin_sub_tab2 = st.tabs([
            "📊 簽到進度總覽與匯出", 
            "👥 會友名單編輯"
        ])

        with admin_sub_tab1:
            st.markdown("### 📊 全會友讀經簽到進度總表")

            time_range = st.selectbox(
                "📅 請選擇匯出與統計時間區間：",
                ["最近 4 週", "第一季 (W01~W13)", "第二季 (W14~W26)", "第三季 (W27~W39)", "第四季 (W40~W52)", "半年 (26 週)", "全年度 (52 週)"]
            )

            if time_range == "第一季 (W01~W13)":
                start_w, end_w = 1, 13
            elif time_range == "第二季 (W14~W26)":
                start_w, end_w = 14, 26
            elif time_range == "第三季 (W27~W39)":
                start_w, end_w = 27, 39
            elif time_range == "第四季 (W40~W52)":
                start_w, end_w = 40, 52
            elif time_range == "最近 4 週":
                start_w, end_w = max(1, current_week_num - 3), current_week_num
            elif time_range == "半年 (26 週)":
                start_w, end_w = max(1, current_week_num - 25), current_week_num
            else:
                start_w, end_w = 1, 52

            df_pivot = generate_pivot_report(PLAN_YEAR, 52)
            target_cols = ["member_name", "完成週數", "完成率"] + [f"Y{PLAN_YEAR}-W{w:02d}" for w in range(start_w, end_w + 1)]
            df_pivot_filtered = df_pivot[target_cols]

            st.dataframe(df_pivot_filtered, use_container_width=True, height=400)

            csv_bytes = df_pivot_filtered.to_csv(index=False, encoding="utf-8-sig").encode("utf-8-sig")

            st.download_button(
                label=f"📥 下載【{time_range}】簽到統計 Excel 報表 (CSV)",
                data=csv_bytes,
                file_name=f"Church_Attendance_{time_range}_{datetime.datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )

            st.divider()
            st.markdown("#### 🛠️ 誤簽撤銷 / 刪除紀錄區")
            col_del1, col_del2, col_del3, col_del4 = st.columns([2, 2, 2, 1.5])

            with col_del1:
                del_member = st.selectbox("選擇要修正的會友：", member_list)
            with col_del2:
                del_year_num = st.number_input("選擇年份：", min_value=1, max_value=4, value=PLAN_YEAR)
            with col_del3:
                del_week_num = st.number_input("選擇週數 (1~52)：", min_value=1, max_value=52, value=current_week_num)
                del_week_key = f"Y{del_year_num}-W{del_week_num:02d}"
            with col_del4:
                st.write("")
                st.write("")
                if st.button("❌ 撤銷此簽到", type="secondary"):
                    delete_single_record(del_week_key, del_member)
                    st.toast(f"已成功刪除 {del_member} 在【{del_week_key}】的紀錄！")
                    st.rerun()

        with admin_sub_tab2:
            st.markdown("### 👥 管理會友名單")
            st.write("可在下方文字框中新增或修改會友姓名（每行一位）：")

            current_m_text = "\n".join(member_list)
            new_m_text = st.text_area("會友名單列表：", value=current_m_text, height=350)

            if st.button("💾 儲存名單變更"):
                updated_names = [name.strip() for name in new_m_text.split("\n") if name.strip()]
                save_members(updated_names)
                st.success("🎉 會友名單更新成功！")
                st.rerun()
